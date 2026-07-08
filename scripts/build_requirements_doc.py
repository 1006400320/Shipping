from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path(__file__).resolve().parents[1] / "物流过程管控系统_需求规格说明书.docx"

FONT_EAST_ASIA = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(85, 85, 85)
LIGHT_GRAY = "F2F4F7"
CALL_OUT = "F4F6F9"
RISK_FILL = "FFF2F2"
OK_FILL = "F0F7F4"

TOTAL_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_EAST_ASIA)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT_LATIN)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_style_font(style, size: float, color: RGBColor | None = None, bold: bool | None = None):
    style.font.name = FONT_LATIN
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_para(doc, text="", style=None, size: float | None = None, bold: bool | None = None, color=None, after=6):
    p = doc.add_paragraph(style=style)
    set_paragraph_spacing(p, after=after)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph_spacing(
        p,
        before={1: 16, 2: 12, 3: 8}.get(level, 8),
        after={1: 8, 2: 6, 3: 4}.get(level, 4),
    )
    run = p.add_run(text)
    set_run_font(
        run,
        size={1: 16, 2: 13, 3: 12}.get(level, 12),
        color=BLUE if level in (1, 2) else DARK_BLUE,
        bold=True,
    )
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=4, line=1.167)
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(p, after=4, line=1.167)
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in CELL_MARGINS.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def clear_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph._element.getparent().remove(paragraph._element)


def cell_text(cell, text, bold=False, fill=None, size=9.5, color=BLACK, align=None):
    clear_cell(cell)
    p = cell.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.12)
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    if fill:
        set_cell_shading(cell, fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell_text(table.rows[0].cells[i], header, bold=True, fill=LIGHT_GRAY, size=font_size)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cell_text(cells[i], value, size=font_size)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=4)
    return table


def add_callout(doc, title, body, fill=CALL_OUT):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TOTAL_WIDTH_DXA])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    clear_cell(cell)
    p = cell.add_paragraph()
    set_paragraph_spacing(p, after=3, line=1.12)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, after=0, line=1.12)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10, color=BLACK)
    set_cell_shading(cell, fill)
    set_cell_margins(cell)
    p3 = doc.add_paragraph()
    set_paragraph_spacing(p3, after=4)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    set_style_font(styles["Normal"], 11, BLACK)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10

    for name, size, color in [
        ("Title", 24, BLACK),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        set_style_font(styles[name], size, color, True)

    for name in ("List Bullet", "List Number"):
        set_style_font(styles[name], 11, BLACK)
        styles[name].paragraph_format.space_after = Pt(4)
        styles[name].paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = "物流过程管控系统需求规格说明书"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header, after=0)
    if header.runs:
        set_run_font(header.runs[0], size=9, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("第 ")
    set_run_font(run, size=9, color=GRAY)
    add_page_number(footer)
    run2 = footer.add_run(" 页")
    set_run_font(run2, size=9, color=GRAY)


def add_cover(doc):
    for _ in range(4):
        add_para(doc, "", after=8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=8)
    r = p.add_run("深圳市捷顺科技实业股份有限公司")
    set_run_font(r, size=15, color=GRAY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=4)
    r = p.add_run("物流过程管控系统")
    set_run_font(r, size=26, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=26)
    r = p.add_run("需求规格说明书")
    set_run_font(r, size=20, color=BLUE, bold=True)

    add_table(
        doc,
        ["项目", "内容"],
        [
            ("文档版本", "V1.0"),
            ("编制日期", "2026-05-20"),
            ("适用范围", "发货作业台、扫码作业、物流费用、基础配置等首版功能"),
            ("文档状态", "草稿"),
            ("保密等级", "内部资料"),
        ],
        [2000, 7360],
        font_size=10,
    )
    doc.add_page_break()


def add_front_matter(doc):
    add_heading(doc, "文档控制", 1)
    add_heading(doc, "修订记录", 2)
    add_table(
        doc,
        ["版本", "日期", "修订说明", "修订人"],
        [("V1.0", "2026-05-20", "根据当前项目代码、页面原型和物流系统方案形成初版需求规格。", "项目组")],
        [1200, 1600, 5060, 1500],
    )
    add_heading(doc, "评审与确认", 2)
    add_table(
        doc,
        ["角色", "姓名", "意见", "日期"],
        [
            ("业务负责人", "", "", ""),
            ("产品负责人", "", "", ""),
            ("研发负责人", "", "", ""),
            ("测试负责人", "", "", ""),
        ],
        [2200, 1800, 3860, 1500],
    )
    add_heading(doc, "目录", 1)
    for item in [
        "1 引言",
        "2 项目概述",
        "3 用户角色与职责",
        "4 业务流程与状态机",
        "5 功能性需求",
        "6 页面与交互需求",
        "7 数据需求",
        "8 接口需求",
        "9 非功能性需求",
        "10 首版交付范围",
        "11 验收标准",
        "12 附录",
    ]:
        add_para(doc, item, after=2)
    doc.add_page_break()


def add_introduction(doc):
    add_heading(doc, "1 引言", 1)
    add_heading(doc, "1.1 编写目的", 2)
    add_para(
        doc,
        "本文档用于明确物流过程管控系统首版建设范围、业务规则、功能需求、数据结构、接口边界、非功能要求和验收标准，为产品评审、研发实现、测试验收和后续迭代提供依据。",
    )
    add_heading(doc, "1.2 项目背景", 2)
    add_para(
        doc,
        "当前项目围绕捷顺发货物流业务，重点解决发货资料完善、打印、扫码拣配、扫码抽检、扫码封箱、DNA 录入、发厂确认、物流单号、签收、对账、改价确认和报销跟踪的端到端闭环。系统现阶段以 Vue 3 + Vite 前端原型为主，已包含发货作业台、费用流程、物流配置和基础档案页面。",
    )
    add_heading(doc, "1.3 文档范围", 2)
    for text in [
        "覆盖发货任务从资料完善到费用报销完成的首版业务闭环。",
        "覆盖扫码枪连续扫码、校验、异常提示和扫码记录留痕。",
        "覆盖物流承运商、运费、配件箱、发货人等基础配置。",
        "覆盖前端页面行为、核心接口规划和数据模型建议。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "1.4 术语与缩写", 2)
    add_table(
        doc,
        ["术语", "说明"],
        [
            ("DNA", "大件产品需要录入的外部系统编号，用于发厂前强校验。"),
            ("扫码枪", "按键盘输入处理的条码采集设备，页面需要支持自动聚焦和回车提交。"),
            ("交货单", "系统内承载收货信息、物料清单、物流信息和费用状态的核心业务单据。"),
            ("对账", "物流公司上传费用数据后，财务进行核价与改价确认的流程。"),
            ("报销闭环", "对账确认后自动发起并跟踪总部财务处理结果的费用闭环。"),
        ],
        [2200, 7160],
    )


def add_overview(doc):
    add_heading(doc, "2 项目概述", 1)
    add_heading(doc, "2.1 系统定位", 2)
    add_callout(
        doc,
        "系统定位",
        "物流过程管控系统是内部发货物流流程管控工具，不定位为通用 WMS 或运输调度系统。首版聚焦单据状态可控、扫码作业可追踪、大件 DNA 强校验、物流费用对账和报销闭环。",
    )
    add_heading(doc, "2.2 核心业务闭环", 2)
    for step in [
        "发货资料完善",
        "打印备料/发货/封箱相关单据",
        "扫码拣配",
        "扫码抽检",
        "扫码封箱",
        "大件产品 DNA 录入",
        "确认发厂",
        "物流单号录入",
        "运输与签收",
        "物流费用对账",
        "改价确认",
        "报销跟踪",
    ]:
        add_number(doc, step)
    add_heading(doc, "2.3 建设目标", 2)
    for text in [
        "形成以交货单为中心的流程状态机，避免跨节点误操作。",
        "将拣配、抽检、封箱等关键动作沉淀为扫码记录，支持追溯和异常处理。",
        "对含大件产品的交货单强制进行 DNA 编号录入和校验。",
        "对物流单号、签收、作废、对账、改价和报销进行统一跟踪。",
        "提供基础配置能力，支持运费规则、配件箱、承运公司、提货人和发货人维护。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "2.4 系统边界", 2)
    add_table(
        doc,
        ["边界类型", "包含", "不包含"],
        [
            ("发货作业", "资料完善、打印、扫码、封箱、发厂确认、签收、作废", "完整库存 WMS、仓位策略、复杂库内调拨"),
            ("物流管理", "承运商信息、物流单号、对账单、改价确认", "实时车辆定位、自动派车、路径优化"),
            ("费用闭环", "费用项、核价、改价、多方确认、报销跟踪", "复杂费用规则引擎、多承运商自动比价"),
            ("系统实现", "前端页面、接口规划、数据模型", "外部 ERP/OA/财务系统的最终集成细节"),
        ],
        [1600, 3980, 3780],
    )


def add_roles(doc):
    add_heading(doc, "3 用户角色与职责", 1)
    add_table(
        doc,
        ["角色", "主要职责", "关键权限"],
        [
            ("发货资料维护人", "完善收货、合同、销售订单、物料清单、成本中心、发货方式等信息。", "新建/编辑交货单、导入物料、打印单据"),
            ("包管员", "使用扫码枪完成物料拣配，处理少扫、多扫、错扫异常。", "拣配扫码、查看拣配清单"),
            ("质量员", "对物料或箱码进行抽检，登记通过或不通过原因。", "抽检扫码、登记抽检结果"),
            ("封箱员", "扫描箱码和物料码，形成箱内明细并打印调拨单。", "封箱扫码、绑定箱码、打印调拨单"),
            ("DNA 录入员", "对大件产品录入 DNA 编号并完成校验。", "录入/校验 DNA"),
            ("物控/装车库", "确认发厂并触发物流发运节点。", "发厂确认、查看装车状态"),
            ("物流公司", "填写物流单号、上传对账单、确认改价。", "物流单号维护、对账单上传、改价确认"),
            ("财务", "核价、改价审核、发起或跟踪报销。", "费用核价、改价确认、报销跟踪"),
            ("工厂/总部财务", "对改价和费用报销进行确认。", "费用确认、报销处理"),
            ("客户/用户", "接收发货通知并提供签收信息。", "签收信息反馈"),
        ],
        [1700, 5060, 2600],
        font_size=9,
    )


def add_process_and_status(doc):
    add_heading(doc, "4 业务流程与状态机", 1)
    add_heading(doc, "4.1 主流程", 2)
    for text in [
        "发货任务创建后进入待完善状态，由资料维护人补全收货信息、成本中心、发运方式、结算方式和物料清单。",
        "资料完整后打印送货单、交货单和封箱相关单据，打印完成后允许进入扫码拣配。",
        "拣配、抽检、封箱依次通过扫码枪完成，系统按条码、数量、归属关系和状态进行校验。",
        "封箱完成后，如交货单包含大件产品，则必须先录入并校验 DNA 编号。",
        "发厂确认后，由物流公司填写物流单号；录入单号后进入运输和签收节点。",
        "签收或产生费用后进入对账，财务核价；若有改价，需物流公司、工厂和总部财务确认后才能报销。",
    ]:
        add_number(doc, text)
    add_heading(doc, "4.2 交货单状态机", 2)
    add_table(
        doc,
        ["状态编码", "状态名称", "允许进入条件", "主要控制规则"],
        [
            ("DRAFT", "待完善", "发货任务创建后", "资料未完整前不可打印、不可拣配"),
            ("WAIT_PRINT", "待打印", "发货资料完整", "打印完成后才能进入拣配"),
            ("WAIT_PICK", "待拣配", "单据已打印", "仅允许包管员扫码拣配"),
            ("PICKING", "拣配中", "包管员开始扫码", "少扫不可进入抽检，多扫/错扫生成异常"),
            ("WAIT_QC", "待抽检", "应拣物料全部完成", "必须由质量员扫码抽检"),
            ("QC_CHECKING", "抽检中", "质量员开始扫码", "不通过必须填写原因并退回或进入异常处理"),
            ("WAIT_PACK", "待封箱", "抽检通过", "未抽检通过不可封箱"),
            ("PACKING", "封箱中", "封箱员开始扫码", "箱码唯一，物料不得重复装箱或跨箱冲突"),
            ("WAIT_DNA", "待录入 DNA", "已封箱且包含大件产品", "DNA 未完成不可确认发厂"),
            ("WAIT_FACTORY_CONFIRM", "待确认发厂", "已封箱且 DNA 完成或无需 DNA", "发厂确认后才能填写物流单号"),
            ("WAIT_WAYBILL", "待物流单号", "已确认发厂", "未填写物流单号不可进入运输中"),
            ("IN_TRANSIT", "运输中", "物流公司已填写物流单号", "支持运输状态和签收信息追踪"),
            ("SIGNED", "已签收", "客户/用户签收完成", "已签收不可直接作废，只能走异常冲销"),
            ("WAIT_RECONCILE", "待对账", "已签收或已产生物流费用", "等待物流公司上传对账单"),
            ("RECONCILING", "对账中", "物流公司已上传对账单", "财务核价，改价需多方确认"),
            ("WAIT_REIMBURSE", "待报销", "对账确认完成", "系统进入报销跟踪"),
            ("REIMBURSED", "已报销", "报销完成", "费用闭环完成"),
            ("VOIDED", "已作废", "作废审批通过", "已作废不可继续发货；如已产生费用，可继续对账"),
            ("FINISHED", "已结束", "签收和费用闭环均完成", "流程关闭，仅允许查询和追溯"),
        ],
        [1700, 1500, 3100, 3060],
        font_size=8.5,
    )
    add_heading(doc, "4.3 状态控制规则", 2)
    for text in [
        "未打印不能拣配；未完成拣配不能抽检；未抽检通过不能封箱。",
        "拣配、抽检、封箱必须通过扫码记录推进，不能只靠人工点击通过。",
        "封箱完成后物料原则上不能直接修改，必须走异常处理流程。",
        "包含大件产品时，未完成 DNA 录入不能确认发厂。",
        "未确认发厂不能填写物流单号；未填写物流单号不能进入运输中。",
        "已作废单据不能继续发货，但如已产生物流费用，允许继续对账和报销。",
    ]:
        add_bullet(doc, text)


def add_functional_requirements(doc):
    add_heading(doc, "5 功能性需求", 1)
    add_table(
        doc,
        ["模块", "需求编号", "需求说明", "优先级"],
        [
            ("发货作业台", "FR-001", "支持按状态页签、关键字、到货日期、承运公司筛选发货任务。", "P0"),
            ("发货作业台", "FR-002", "发货任务列表展示送货单号、状态、出货申请单号、交货单号、调拨单号、到货日、合同号、销售单号、收货单位和可执行操作。", "P0"),
            ("资料完善", "FR-003", "支持维护收货单位、成本中心、发运方式、结算方式、始发地、目的地、详细地址、交付说明、合同号、销售订单号、收货人和电话。", "P0"),
            ("资料完善", "FR-004", "支持导入物料清单，字段包括物料编号、描述、单位、实发数、件数、运价、运价合计和备注。", "P0"),
            ("资料完善", "FR-005", "支持承运方签收信息、发货方信息和收货方签收信息维护。", "P1"),
            ("打印", "FR-006", "支持打印送货单、交货单、调拨单，并在打印后推进状态。", "P0"),
            ("扫码拣配", "FR-007", "扫码对象包括交货单号、物料码、SKU 码和预装箱箱码。", "P0"),
            ("扫码拣配", "FR-008", "系统校验条码归属、计划数量、重复扫描和占用状态，少扫不可进入抽检。", "P0"),
            ("扫码抽检", "FR-009", "质量员扫描物料或箱码后登记抽检结果，不通过必须填写原因。", "P0"),
            ("扫码封箱", "FR-010", "支持箱码绑定、箱内物料扫描、封箱进度、重复装箱拦截和调拨单码绑定。", "P0"),
            ("DNA 录入", "FR-011", "交货单包含大件产品时必须录入 DNA 编号，校验完成后才能确认发厂。", "P0"),
            ("发厂与物流", "FR-012", "发厂确认后由物流公司填写承运信息和物流单号，未填写前不能进入运输中。", "P0"),
            ("签收与作废", "FR-013", "支持签收状态登记和作废申请；已签收不允许直接作废。", "P1"),
            ("对账核价", "FR-014", "物流公司上传对账单后，财务核价并维护基础运费、送货费、中转费、延期扣费和其他费用。", "P0"),
            ("改价确认", "FR-015", "存在改价时，必须记录改价前后金额、原因、发起人、确认人、确认时间和附件。", "P0"),
            ("报销跟踪", "FR-016", "对账确认完成后自动进入报销流程并跟踪总部财务处理状态。", "P0"),
            ("物流配置", "FR-017", "支持维护运费规则、配件箱、承运公司、提货人和发货人，并支持查询、启停、新增和编辑。", "P0"),
            ("基础档案", "FR-018", "支持维护物料档案、大件标识和承运商档案。", "P1"),
            ("操作日志", "FR-019", "所有扫码、状态推进、费用改价、作废和配置变更必须记录操作人、时间和结果。", "P0"),
        ],
        [1450, 1050, 5860, 1000],
        font_size=8.5,
    )
    add_heading(doc, "5.1 扫码通用要求", 2)
    for text in [
        "页面进入扫码作业时，扫码输入框自动聚焦。",
        "扫码枪按键盘输入处理，扫码后通过回车或自动提交完成业务校验。",
        "支持连续扫描，成功后选中输入内容，方便下一次扫码。",
        "扫码成功、失败、异常必须在页面上给出明确提示。",
        "每次扫码必须写入扫码记录，包含业务类型、条码、设备号、操作人、结果和失败原因。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "5.2 物流配置要求", 2)
    add_table(
        doc,
        ["配置项", "字段", "规则"],
        [
            ("运费", "承运公司、类型、状态、出发省市、目的省市、运输费、时效、更新时间、更新人", "可查询、导入、导出；运输费不得为负数。"),
            ("配件箱", "型号、描述、状态、长、宽、高、体积、更新时间、更新人", "体积按长宽高自动计算；尺寸不得小于 0。"),
            ("承运公司", "公司名称、类型、重抛比、货运委托书、状态、提货人", "支持提货人管理；停用承运公司不可用于发货。"),
            ("发货人", "姓名、工号、电话、备注、状态", "仅正常状态可被发货完善页选择。"),
        ],
        [1600, 5200, 2560],
        font_size=8.8,
    )


def add_pages(doc):
    add_heading(doc, "6 页面与交互需求", 1)
    add_table(
        doc,
        ["页面", "主要功能", "关键交互"],
        [
            ("发货作业台", "发货任务列表、状态页签、查询、重置、导出、任务详情、节点概览。", "点击任务号进入详情；待完善任务可进入完善页。"),
            ("完善页", "维护送货单、成本中心、发运方式、结算方式、物料清单、费用、承运方和签收信息。", "成本中心弹窗选择；承运公司联动提货人；费用自动合计。"),
            ("扫码拣配页", "连续扫描物料码，展示计划、已拣和状态。", "多扫、错扫进入异常处理。"),
            ("扫码抽检页", "扫描物料或箱码并登记抽检结果。", "不通过时退回拣配或进入异常。"),
            ("扫码封箱页", "扫描箱码和物料码，展示装箱进度、箱清单和操作痕迹。", "输入 MAT-L3002 可模拟补扫成功；重复装箱提示异常。"),
            ("DNA 录入页", "展示大件清单，录入 DNA 编号并校验。", "未完成 DNA 时发厂确认锁定。"),
            ("物流单号页", "填写物流公司和物流单号，展示运输节点。", "发厂确认后允许录入单号。"),
            ("对账核价页", "展示费用项、预估费用和对账状态。", "支持费用核价和改价确认。"),
            ("报销跟踪页", "展示对账、核价、发起报销、总部确认、已报销流程。", "对账完成后自动推进。"),
            ("物流配置页", "维护运费、配件箱、承运公司、发货人。", "支持状态开关、新增、编辑、导入、导出和本地持久化。"),
            ("物料档案页", "维护物料名称、大件标识和状态。", "大件标识触发 DNA 强校验。"),
            ("物流公司页", "维护承运商、联系人和启停状态。", "供物流单号和对账流程使用。"),
        ],
        [1700, 4960, 2700],
        font_size=8.5,
    )
    add_heading(doc, "6.1 通用交互规范", 2)
    for text in [
        "查询区域应支持条件重置，并保留清晰的查询结果提示。",
        "状态类字段应使用明确文案和颜色区分正常、待处理、异常。",
        "操作按钮应按当前状态动态显示，禁止无效节点操作。",
        "表格空数据时应显示空状态文案。",
        "配置类数据保存失败时，应给出可见错误提示，不得静默失败。",
    ]:
        add_bullet(doc, text)


def add_data_requirements(doc):
    add_heading(doc, "7 数据需求", 1)
    add_heading(doc, "7.1 核心数据实体", 2)
    add_table(
        doc,
        ["实体", "说明", "关键字段"],
        [
            ("shipment_order", "交货单主表", "shipment_no、customer_name、receiver_name、status、has_large_item、dna_required、waybill_no、signed_at、created_at"),
            ("shipment_item", "发货物料明细", "shipment_id、sku_code、material_code、item_name、planned_qty、picked_qty、packed_qty、is_large_item"),
            ("scan_record", "扫码记录", "shipment_no、business_type、scan_code、scan_code_type、operator_id、device_no、scan_result、fail_reason、created_at"),
            ("package_box", "箱信息", "box_no、shipment_id、status、sealed_by、sealed_at"),
            ("package_box_item", "箱内物料", "box_id、shipment_item_id、sku_code、material_code、planned_qty、actual_qty"),
            ("dna_record", "DNA 录入记录", "shipment_id、material_code、box_no、dna_no、verifier、verify_status、verified_at"),
            ("logistics_reconciliation", "物流对账", "shipment_id、logistics_company_id、statement_file_id、base_fee、delivery_fee、total_fee、adjusted_total_fee、status"),
            ("freight_config", "运费配置", "carrier、type、fromProvince、fromCity、toProvince、toCity、volumeFee、weightFee、leadDays、status"),
            ("package_box_config", "配件箱配置", "model、description、length、width、height、volume、status"),
            ("carrier_config", "承运公司配置", "carrier、type、weightRatio、freightAuthorizationLetter、pickupPeople、status"),
            ("sender_config", "发货人配置", "sender、employeeNo、phone、remark、status"),
        ],
        [2300, 2200, 4860],
        font_size=8.2,
    )
    add_heading(doc, "7.2 数据规则", 2)
    for text in [
        "交货单号、箱码、物流单号、DNA 编号在对应业务范围内必须唯一。",
        "扫码记录不得物理删除，异常扫描也必须保留失败原因。",
        "费用改价必须保留改价前金额、改价后金额、原因、确认链路和附件。",
        "停用的承运公司、提货人、发货人和配置项不得被新业务单据选择。",
        "本地原型使用 IndexedDB 保存物流配置，后端实现时应替换为服务端持久化接口。",
    ]:
        add_bullet(doc, text)


def add_api_requirements(doc):
    add_heading(doc, "8 接口需求", 1)
    add_heading(doc, "8.1 扫码接口", 2)
    add_table(
        doc,
        ["接口", "方法", "说明", "关键入参", "关键出参"],
        [
            ("/api/scan", "POST", "统一扫码入口", "shipmentNo、businessType、scanCode、deviceNo", "success、message、nextStatus、progress"),
            ("/api/shipments/{shipmentNo}/actions/{action}", "POST", "状态推进", "action、operator、remark", "success、status、message"),
            ("/api/logistics/waybills", "POST", "录入物流单号", "shipmentNo、carrierId、waybillNo", "success、nextStatus"),
            ("/api/logistics/reconciliations", "POST", "上传对账单", "shipmentNo、feeItems、statementFileId", "success、reconciliationId"),
            ("/api/logistics/reconciliations/{id}/confirm-price-change", "POST", "确认改价", "confirmRole、confirmUser、result、remark", "success、nextConfirmRole、status"),
        ],
        [2500, 900, 1900, 2500, 1560],
        font_size=7.8,
    )
    add_heading(doc, "8.2 状态推进动作", 2)
    add_table(
        doc,
        ["Action", "触发节点", "前置条件", "后置状态"],
        [
            ("confirm-print", "打印完成", "资料完整且单据可打印", "WAIT_PICK"),
            ("complete-pick", "拣配完成", "应拣物料全部扫码完成", "WAIT_QC"),
            ("complete-qc", "抽检完成", "抽检结果全部通过", "WAIT_PACK"),
            ("complete-pack", "封箱完成", "物料全部装箱且箱码有效", "WAIT_DNA 或 WAIT_FACTORY_CONFIRM"),
            ("complete-dna", "DNA 完成", "大件 DNA 全部录入校验通过", "WAIT_FACTORY_CONFIRM"),
            ("confirm-factory", "发厂确认", "封箱完成且无需或已完成 DNA", "WAIT_WAYBILL"),
            ("bind-waybill", "录入物流单号", "已确认发厂", "IN_TRANSIT"),
            ("sign", "签收", "运输中且签收信息有效", "SIGNED"),
            ("void", "作废", "满足作废审批条件", "VOIDED"),
            ("submit-reconciliation", "上传对账单", "已签收或已产生费用", "RECONCILING"),
            ("confirm-reconciliation", "对账确认", "核价完成且改价确认完成", "WAIT_REIMBURSE"),
            ("complete-reimburse", "报销完成", "总部财务处理完成", "REIMBURSED 或 FINISHED"),
        ],
        [2100, 1800, 3660, 1800],
        font_size=8.2,
    )


def add_nonfunctional(doc):
    add_heading(doc, "9 非功能性需求", 1)
    add_table(
        doc,
        ["类别", "需求"],
        [
            ("性能", "常规列表查询响应时间不超过 2 秒；扫码提交从回车到反馈不超过 1 秒。"),
            ("可用性", "扫码页面必须支持自动聚焦、连续扫描、明确成功/失败反馈和异常原因展示。"),
            ("可靠性", "状态推进接口必须幂等；重复提交不得造成重复装箱、重复报销或重复状态推进。"),
            ("安全", "按角色限制页面和操作权限；费用、作废、改价、配置变更等敏感操作必须审计。"),
            ("数据一致性", "交货单状态、扫码记录、箱明细、DNA 记录、物流费用需在同一业务事务内保持一致。"),
            ("兼容性", "前端需支持现代 Chromium 内核浏览器；扫码枪按键盘输入兼容。"),
            ("可维护性", "状态编码、角色权限、费用项和配置项需集中管理，避免散落在页面逻辑中。"),
            ("可追溯性", "关键业务动作必须保留操作人、时间、来源设备、前后状态和失败原因。"),
        ],
        [1800, 7560],
        font_size=9,
    )


def add_scope_and_acceptance(doc):
    add_heading(doc, "10 首版交付范围", 1)
    add_heading(doc, "10.1 必须包含", 2)
    for text in [
        "交货单创建、资料完善、打印标记和任务查询。",
        "扫码拣配、扫码抽检、扫码封箱和异常提示。",
        "箱码与箱内物料明细维护。",
        "大件产品 DNA 录入与发厂前强校验。",
        "发厂确认、物流单号录入、签收和作废。",
        "对账单上传、财务核价、改价确认、报销状态跟踪。",
        "运费、配件箱、承运公司、提货人和发货人配置。",
        "操作日志和扫码记录留痕。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "10.2 暂不包含", 2)
    for text in [
        "实时车辆定位。",
        "路线优化和自动派车。",
        "完整库存 WMS。",
        "复杂费用规则引擎。",
        "多承运商自动比价。",
        "外部 ERP、OA、财务系统的正式联调和上线切换方案。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "11 验收标准", 1)
    add_table(
        doc,
        ["场景", "前置条件", "操作", "预期结果"],
        [
            ("正常发货闭环", "发货资料完整且物料齐套", "依次完成打印、拣配、抽检、封箱、DNA、发厂、物流单号、签收、对账、报销", "流程进入已结束，扫码和费用记录完整"),
            ("拣配少扫", "交货单处于待拣配", "少扫部分物料后尝试进入抽检", "系统禁止进入抽检并提示缺失物料"),
            ("拣配多扫", "物料计划数量已满足", "继续扫描同一物料", "系统拦截并记录失败原因"),
            ("扫错物料", "当前交货单不包含该物料", "扫描错误物料码", "系统提示不属于当前交货单"),
            ("抽检不通过", "拣配完成", "登记抽检不通过并填写原因", "退回拣配或进入异常处理"),
            ("未抽检封箱", "交货单未完成抽检", "尝试进入封箱", "系统禁止封箱"),
            ("箱码重复", "箱码已被绑定", "再次绑定同一箱码", "系统禁止绑定并记录异常"),
            ("大件未录 DNA", "交货单包含大件产品", "未录入 DNA 时确认发厂", "系统禁止确认发厂"),
            ("未填物流单号", "已确认发厂", "不填单号尝试进入运输中", "系统禁止进入运输中"),
            ("已作废继续发货", "交货单已作废", "尝试继续拣配或发厂", "系统禁止继续发货"),
            ("已作废但有费用", "已产生物流费用", "进入费用对账", "允许继续费用对账和报销"),
            ("对账无改价", "物流公司上传对账单", "财务核价无改价", "系统自动进入报销流程"),
            ("对账有改价", "财务核价发现改价", "物流、工厂、总部财务依次确认", "确认完成后进入报销流程"),
        ],
        [1600, 2600, 3000, 2160],
        font_size=7.8,
    )


def add_appendix(doc):
    add_heading(doc, "12 附录", 1)
    add_heading(doc, "12.1 当前项目技术信息", 2)
    add_table(
        doc,
        ["项目项", "说明"],
        [
            ("前端框架", "Vue 3.5.14"),
            ("构建工具", "Vite 6.3.5"),
            ("本地存储", "IndexedDB，当前用于物流配置原型数据保存"),
            ("运行脚本", "npm run dev、npm run build、npm run preview"),
            ("主要页面", "WorkbenchPage、CompleteDeliveryPage、WaybillPage、ReconcilePage、ReimbursePage、FreightConfigPage、MaterialPage、CarrierPage"),
        ],
        [2200, 7160],
    )
    add_heading(doc, "12.2 风险与待确认事项", 2)
    add_table(
        doc,
        ["事项", "风险", "建议处理"],
        [
            ("参考格式文件为空", "无法复用原始 DOCX 的页眉、目录、样式和模板控件。", "业务方补充正式模板后，可按模板二次套版。"),
            ("外部系统集成", "ERP/OA/财务/DNA 系统接口字段和流程未最终确认。", "在接口评审阶段补充系统间字段映射和异常补偿机制。"),
            ("权限模型", "当前前端原型未落地完整权限控制。", "后端设计阶段补充角色、组织、数据范围和审批权限。"),
            ("扫码设备差异", "不同扫码枪回车符、输入速度和编码可能不同。", "测试阶段覆盖主流设备并提供配置项。"),
        ],
        [2200, 2860, 4300],
        font_size=8.8,
    )


def build():
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_introduction(doc)
    add_overview(doc)
    add_roles(doc)
    add_process_and_status(doc)
    add_functional_requirements(doc)
    add_pages(doc)
    add_data_requirements(doc)
    add_api_requirements(doc)
    add_nonfunctional(doc)
    add_scope_and_acceptance(doc)
    add_appendix(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
